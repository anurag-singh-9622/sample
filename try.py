```python
from docx import Document  # Importing the Document class from the docx module
from bs4 import BeautifulSoup  # Importing BeautifulSoup for parsing HTML content

# HTML content (replace with your actual HTML)
html_content = '''
<!DOCTYPE html>
<html>
<head>
    <title>Sample HTML</title>
    <style>
        h1 { color: blue; }
        p { font-size: 14px; }
    </style>
</head>
<body>
    <h1>Hello, Markdown!</h1>
    <p>This is a sample paragraph.</p>
</body>
</html>
'''

# Create a new Word document
doc = Document()  # creating an instance of the Document class

# Parse the HTML content using BeautifulSoup
soup = BeautifulSoup(html_content, 'html.parser')  # parsing the HTML content using BeautifulSoup

# Extract and add text and formatting to the document
for tag in soup.find_all(['h1', 'p']):  # iterating through 'h1' and 'p' tags in the HTML content
    if tag.name == 'h1':  # checking if the tag is 'h1'
        doc.add_heading(tag.text, level=1)  # adding a heading with the text from the 'h1' tag
    elif tag.name == 'p':  # checking if the tag is 'p'
        doc.add_paragraph(tag.text)  # adding a paragraph with the text from the 'p' tag

# Save the document in memory (in a variable)
docx_output = doc  # storing the document in a variable

# Print the first paragraph (for demonstration purposes)
print(docx_output.paragraphs[0].text)  # printing the text of the first paragraph
```